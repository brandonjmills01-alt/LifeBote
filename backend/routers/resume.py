"""
routers/resume.py
Features 1 & 3: Resume Tailoring + Optimization Score
Feature 2/3 support (T3 Auto Apply): Resume profile extraction
T2: Resume upload (PDF + DOCX) + Export (PDF + DOCX)

Endpoints:
    POST /api/resume/tailor          — AI rewrite + scoring (existing)
    POST /api/resume/upload          — extract text from PDF or DOCX (new)
    POST /api/resume/upload-pdf      — legacy alias for PDF-only upload
    POST /api/resume/analyze         — extract profile (skills, seniority, etc.)
    POST /api/resume/export/pdf      — render resume text to a downloadable PDF
    POST /api/resume/export/docx     — render resume text to a downloadable DOCX
"""

import io
import re
from datetime import datetime, timezone

from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse

from models.schemas import (
    ResumeTailorRequest, ResumeTailorResponse, ResumeScore,
    ResumeUploadResponse, ResumeAnalyzeRequest, ResumeProfile,
    ResumeExportRequest,
)
from services import ai_service

router = APIRouter()


# ── Tailor (existing) ────────────────────────────────────────────────────────

@router.post("/tailor", response_model=ResumeTailorResponse)
async def tailor_resume(body: ResumeTailorRequest):
    try:
        r = await ai_service.tailor_resume(
            body.resume_text, body.job_description, body.tone, body.target_role or ""
        )
        return ResumeTailorResponse(
            tailored_resume=r["tailored_resume"],
            changes_summary=r.get("changes_summary", []),
            resume_versions=r.get("resume_versions", {}),
            score=ResumeScore(
                ats_score=r["ats_score"],
                keyword_score=r["keyword_score"],
                readability_score=r["readability_score"],
                overall_score=r["overall_score"],
                missing_keywords=r.get("missing_keywords", []),
                suggestions=r.get("suggestions", []),
                strengths=r.get("strengths", []),
            ),
        )
    except Exception as e:
        # Keep the message friendly; full trace lives in server logs
        raise HTTPException(status_code=500, detail="Could not tailor resume right now. Please try again.")


# ── Upload (PDF or DOCX) ─────────────────────────────────────────────────────

# Max upload size — refuses huge files before reading them into memory
_MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB


def _extract_pdf_text(contents: bytes) -> tuple[str, int]:
    """Return (text, page_count). Raises HTTPException on bad PDFs."""
    import pdfplumber
    try:
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            pages = pdf.pages
            page_count = len(pages)
            text = "\n".join((p.extract_text() or "") for p in pages)
    except Exception as e:
        msg = str(e).lower()
        if "password" in msg or "encrypted" in msg:
            raise HTTPException(status_code=422, detail="This PDF is password-protected. Please remove the password and try again.")
        raise HTTPException(status_code=422, detail="Could not read this PDF. The file may be corrupted.")
    if not text.strip():
        raise HTTPException(
            status_code=422,
            detail="No text could be extracted. This is likely an image-only PDF — try saving it as a text-based PDF or paste your resume manually.",
        )
    return text, page_count


def _extract_docx_text(contents: bytes) -> str:
    """Return text from a DOCX. Raises HTTPException on bad files."""
    from docx import Document
    try:
        doc = Document(io.BytesIO(contents))
    except Exception:
        raise HTTPException(status_code=422, detail="Could not read this DOCX file. It may be corrupted or in an unsupported format.")

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    # Tables often hold contact info / two-column layouts
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text not in parts:
                    parts.append(cell.text)
    text = "\n".join(parts)
    if not text.strip():
        raise HTTPException(status_code=422, detail="No text could be extracted from this DOCX.")
    return text


@router.post("/upload", response_model=ResumeUploadResponse)
async def upload_resume(file: UploadFile = File(...)):
    """
    Accept a PDF or DOCX and return the extracted plain text.
    Errors are friendly so the UI can show them directly to the user.
    """
    name = (file.filename or "").lower()
    if not name:
        raise HTTPException(status_code=400, detail="No file was uploaded.")

    if name.endswith(".pdf"):
        fmt = "pdf"
    elif name.endswith(".docx"):
        fmt = "docx"
    elif name.endswith(".doc"):
        raise HTTPException(status_code=400, detail="Legacy .doc files are not supported. Please save as .docx or PDF and try again.")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF or DOCX.")

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")
    if len(contents) > _MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is too large. Max upload is 10 MB.")

    pages: int | None = None
    if fmt == "pdf":
        text, pages = _extract_pdf_text(contents)
    else:
        text = _extract_docx_text(contents)

    return ResumeUploadResponse(
        filename=file.filename or "",
        size_bytes=len(contents),
        format=fmt,
        text=text,
        pages=pages,
    )


@router.post("/upload-pdf")
async def upload_pdf_legacy(file: UploadFile = File(...)):
    """Legacy PDF-only endpoint. Prefer /api/resume/upload."""
    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted at this endpoint. Use /api/resume/upload for DOCX.")
    contents = await file.read()
    text, page_count = _extract_pdf_text(contents)
    return {"text": text, "pages": page_count}


# ── Analyze: extract a structured profile from resume text ───────────────────

@router.post("/analyze", response_model=ResumeProfile)
async def analyze_resume(body: ResumeAnalyzeRequest):
    """
    Extract a structured profile from raw resume text.
    Used by the Auto Apply tab to drive automatic job matching.
    """
    if not body.resume_text or not body.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text is required.")
    try:
        profile = await ai_service.analyze_resume(body.resume_text)
        return ResumeProfile(**profile)
    except HTTPException:
        raise
    except Exception:
        # Fall back to a minimal profile so the UI never gets stuck
        return ResumeProfile(
            skills=[], experience_years=0, job_titles=[],
            seniority="mid", industries=[], locations=[],
        )


# ── Export: render a tailored resume to a downloadable file ──────────────────

def _split_sections(resume_text: str) -> list[tuple[str | None, list[str]]]:
    """
    Split resume text into (heading, lines) sections.
    Heuristic: a heading is a short all-caps or title-case line that ends without
    a sentence-terminating period. Falls back to a single unnamed section.
    """
    sections: list[tuple[str | None, list[str]]] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush():
        if current_heading is not None or current_lines:
            sections.append((current_heading, list(current_lines)))

    heading_re = re.compile(r"^[A-Z][A-Z &/]{2,40}$")
    for raw in resume_text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            current_lines.append("")
            continue
        is_heading = (
            heading_re.match(stripped)
            or (len(stripped) <= 40 and stripped.endswith(":"))
        )
        if is_heading:
            flush()
            current_heading = stripped.rstrip(":")
            current_lines = []
        else:
            current_lines.append(stripped)
    flush()
    return sections


@router.post("/export/pdf")
async def export_resume_pdf(body: ResumeExportRequest):
    """Render the resume text to a downloadable PDF using reportlab."""
    if not body.resume_text or not body.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text is required.")

    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.colors import HexColor

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
        title=f"Resume — {body.full_name or ''}".strip(),
    )

    styles = getSampleStyleSheet()
    navy = HexColor("#0A2342")
    green = HexColor("#00A86B")

    name_style = ParagraphStyle(
        "Name", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=20, leading=24, textColor=navy, alignment=TA_LEFT, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10, leading=13, textColor=HexColor("#374151"), spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "Heading", parent=styles["Heading2"], fontName="Helvetica-Bold",
        fontSize=12, leading=15, textColor=navy, spaceBefore=10, spaceAfter=4,
        borderPadding=0, leftIndent=0,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10.5, leading=14, textColor=HexColor("#1F2937"), spaceAfter=2,
    )

    story = []
    if body.full_name:
        story.append(Paragraph(body.full_name, name_style))
        # Thin green underline via a single-row table would be heavy; use spacer + contact
    contact_bits = [b for b in [body.contact_email, body.contact_phone, body.contact_linkedin] if b]
    if contact_bits:
        story.append(Paragraph(" · ".join(contact_bits), contact_style))
    else:
        story.append(Spacer(1, 6))

    # Decorative green rule under header (drawn as a 1pt horizontal line via Paragraph)
    story.append(Paragraph('<para><font color="#00A86B"><b>━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━</b></font></para>', body_style))
    story.append(Spacer(1, 6))

    def _escape(t: str) -> str:
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    sections = _split_sections(body.resume_text)
    for heading, lines in sections:
        if heading:
            story.append(Paragraph(_escape(heading), heading_style))
        for line in lines:
            if not line:
                story.append(Spacer(1, 4))
                continue
            # Render bullets visually if line starts with "-", "*", or "•"
            if line.lstrip().startswith(("-", "*", "•")):
                bullet = line.lstrip().lstrip("-*• ").strip()
                story.append(Paragraph(f"• {_escape(bullet)}", body_style))
            else:
                story.append(Paragraph(_escape(line), body_style))

    try:
        doc.build(story)
    except Exception:
        raise HTTPException(status_code=500, detail="Could not generate PDF. Please try again.")

    buf.seek(0)
    fname = _sanitize_filename(body.full_name or "resume") + ".pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.post("/export/docx")
async def export_resume_docx(body: ResumeExportRequest):
    """Render the resume text to a downloadable DOCX using python-docx."""
    if not body.resume_text or not body.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text is required.")

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Standard 0.75" margins on letter paper
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    navy = RGBColor(0x0A, 0x23, 0x42)
    green = RGBColor(0x00, 0xA8, 0x6B)
    body_color = RGBColor(0x1F, 0x29, 0x37)

    # Default style — DM Sans isn't a Word default; Calibri reads cleanly in its place
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = body_color

    if body.full_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(body.full_name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = navy

    contact_bits = [b for b in [body.contact_email, body.contact_phone, body.contact_linkedin] if b]
    if contact_bits:
        p = doc.add_paragraph()
        run = p.add_run(" · ".join(contact_bits))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # Green separator
    sep = doc.add_paragraph()
    sep_run = sep.add_run("─" * 60)
    sep_run.font.color.rgb = green

    for heading, lines in _split_sections(body.resume_text):
        if heading:
            p = doc.add_paragraph()
            run = p.add_run(heading.upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = navy
        for line in lines:
            if not line:
                doc.add_paragraph()
                continue
            stripped = line.lstrip()
            if stripped.startswith(("-", "*", "•")):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(stripped.lstrip("-*• ").strip())
            else:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    try:
        doc.save(buf)
    except Exception:
        raise HTTPException(status_code=500, detail="Could not generate DOCX. Please try again.")
    buf.seek(0)
    fname = _sanitize_filename(body.full_name or "resume") + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


_SAFE_FNAME_RE = re.compile(r"[^A-Za-z0-9_\-]+")


def _sanitize_filename(name: str) -> str:
    cleaned = _SAFE_FNAME_RE.sub("-", name.strip()).strip("-")
    return cleaned or "resume"
